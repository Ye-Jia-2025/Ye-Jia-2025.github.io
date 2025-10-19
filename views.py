from django.http import HttpResponse  # 新增此行
from django.db.models import Max, F
from django.db.models.functions import Coalesce
from django.http import JsonResponse, HttpResponseBadRequest, HttpResponse  # 正确导入
from .models import Conversation, ChatMessage, CourseUpload
from django.utils import timezone
# ... 其他保持不变的导入 ...
import json, logging
from django.utils.timezone import now
import re
from collections import Counter
from pathlib import Path
from typing import List
from .models import Payment
from wechatpayv3 import WeChatPayType
import uuid, io, base64, qrcode
from .pay_client import wxpay
import pdb
from django.db.models import Max, F
from django.db.models.functions import Coalesce
from django.http import JsonResponse, HttpResponseBadRequest
from .models import Conversation, ChatMessage, CourseUpload
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
import os
from django.conf import settings
import pdfplumber
import requests
from django.db.models import Q
from docx import Document
import openai
import faiss
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.docstore.in_memory import InMemoryDocstore
import numpy as np
import threading
import base64
from .sentences_fluency_chi import ai_detect_chi
import json
from django.contrib.auth.models import User
from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.generic import TemplateView
from django.contrib.auth import get_user_model
from .utils import send_sms, generate_code
import shutil
from openai import OpenAI
from datetime import datetime
from .models import LocalKnowledge
from docx import Document
from .my_deep_research import my_deep_thinking, save_markdown_and_pdf
from .my_notebook import my_notebook, generate_markdown
import time  # 新增此行
from django.views.decorators.http import require_http_methods  # 新增此行

# 订阅套餐与价格（单位：分）
PLAN_PRICES = {
    "month": 1,   
    "year": 5, 
}

PLAN_NAMES = {
    "month": "月度订阅",
    "year":  "年度订阅",
}




os.environ["OPENAI_API_BASE"] = "https://api.openai-proxy.org/v1"
os.environ["OPENAI_API_KEY"] = "sk-dA85I4Bo06hodLvVga1VAnOLSKBNGU9InUa4xD7uS8RiTvna"

client = OpenAI(
    api_key='sk-dA85I4Bo06hodLvVga1VAnOLSKBNGU9InUa4xD7uS8RiTvna',
    base_url='https://api.openai-proxy.org/v1'
)

def convert_messages(messages_data):
    api_messages = []
    for msg in messages_data:
        sender = msg.get('sender', '').strip().lower()
        # 判断 sender 包含 user 或 bot（可以根据需要扩展判断逻辑）
        if 'bot' in sender:
            role = "assistant"
        else:
            role = "user"
        content = msg.get('real_message', '')
        api_messages.append({"role": role, "content": content})
    return api_messages
# myapp/views.py 中需存在类似定义
#def wurenji(request):
#    return HttpResponse("This is the wurenji page.")    #2025.03.30/15:28增加
# views.py 修正后的视图函数（使用模板渲染）


from django.shortcuts import render

def chinese_knowledge(request):
    return render(request, 'chinese_knowledge.html')

def biology_ai(request):
    return render(request, 'biology_ai.html')

def biology_knowledge(request):
    return render(request, 'biology_knowledge.html')

def chat_page(request):
    return render(request, 'chat_page.html')

def chemistry_ai(request):
    return render(request, 'chemistry_ai.html')

def chemistry_knowledge(request):
    return render(request, 'chemistry_knowledge.html')

def chinese_ai(request):
    return render(request, 'chinese_ai.html')

def english_ai(request):
    return render(request, 'english_ai.html')

def english_knowledge(request):
    return render(request, 'english_knowledge.html')

def geography_ai(request):
    return render(request, 'geography_ai.html')

def geography_knowledge(request):
    return render(request, 'geography_knowledge.html')

def history_ai(request):
    return render(request, 'history_ai.html')

def history_knowledge(request):
    return render(request, 'history_knowledge.html')

def homework_correction(request):
    return render(request, 'homework_correction.html')

def homework_qa(request):
    return render(request, 'homework_qa.html')

def lab_guidance(request):
    return render(request, 'lab_guidance.html')


def wenzi(request):
    return render(request, 'wenzi.html')#0706见证奇迹


def math_ai(request):
    return render(request, 'math_ai.html')



def YJ(request):
    return render(request, 'YJ.html')




def math_knowledge(request):
    return render(request, 'math_knowledge.html')



def keyan(request):
    return render(request, 'keyan.html')  #0706 update

def physics_ai(request):
    return render(request, 'physics_ai.html')

def physics_knowledge(request):
    return render(request, 'physics_knowledge.html')

def science_animation(request):
    return render(request, 'science_animation.html')



#打烊！
def index(request):
    return render(request, 'index.html')
#25.04.05/14.37新增


#见证奇迹
def index2(request):
    return render(request, 'index2.html')
#25.04.15/08.23新增
def anf(request):
    return render(request, 'anf.html')

def city(request):
    return render(request, 'city.html')

def flfd(request):
    return render(request, 'flfd.html')

def ggkz(request):
    return render(request, 'ggkz.html')

def hupo(request):
    return render(request, 'hupo.html')

def jc(request):
    return render(request, 'jc.html')

def jt(request):
    return render(request, 'jt.html')

def ky(request):
    return render(request, 'ky.html')

def lvhua(request):
    return render(request, 'lvhua.html')

def ly(request):
    return render(request, 'ly.html')

def nongye(request):
    return render(request, 'nongye.html')

def wuliu(request):
    return render(request, 'wuliu.html')

def xhd(request):
    return render(request, 'xhd.html')

def ycjy(request):
    return render(request, 'ycjy.html')

#25.04.16.07.46见证奇迹











def get_response_with_img(system_prompt, my_prompt, history,  image=None):
    content_user=[{"type": "text","text": my_prompt,}]
    if image:
        with open(image[0], "rb") as image_file:
            base64_image = base64.b64encode(image_file.read()).decode('utf-8')
        content_user.append({
            "type": "image_url",
            "image_url": {
                "url": f"data:image/png;base64,{base64_image}"
            },
        })

    response = client.chat.completions.create(
        model="gpt-4o",
        temperature=0.7,
        messages=[
            {"role": "system", "content": system_prompt},
            *history,
            {
                "role": "user",
                "content": content_user,
            }
        ],
    )


    res = response.choices[0].message.content
    return res

def get_response_with_img_stream(system_prompt, my_prompt, history, image=None):
    """
    流式版本的get_response_with_img，返回生成器用于实时token输出
    """
    content_user=[{"type": "text","text": my_prompt,}]
    if image and len(image) > 0:
        with open(image[0], "rb") as image_file:
            base64_image = base64.b64encode(image_file.read()).decode('utf-8')
        content_user.append({
            "type": "image_url",
            "image_url": {
                "url": f"data:image/png;base64,{base64_image}"
            },
        })

    stream = client.chat.completions.create(
        model="gpt-4o",
        temperature=0.7,
        messages=[
            {"role": "system", "content": system_prompt},
            *history,
            {
                "role": "user",
                "content": content_user,
            }
        ],
        stream=True  # 启用流式输出
    )

    for chunk in stream:
        if len(chunk.choices) > 0 and chunk.choices[0].delta.content is not None:
            token = chunk.choices[0].delta.content
            yield token

class ProtectedView(LoginRequiredMixin, TemplateView):
    login_url = '/login/'
    template_name = 'protected_page.html'

CustomUser = get_user_model()  # 等价于 from .models import CustomUser


def read_word_file(file_path):
    """
    读取指定路径的Word文件，并返回整个文档的文本内容
    """
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

@login_required(login_url='/login/')
def your_protected_view(request):
    # 只有登录用户才能访问
    return render(request, 'protected_page.html')

def send_sms_code(request):
    """单独处理验证码发送请求"""
    if request.method == 'POST':
        phone = request.POST.get('phone_number')
        if not phone:
            return JsonResponse({"message": "请输入手机号"}, status=400)

        # 生成6位验证码
        code = generate_code(6)

        try:
            response = send_sms(phone, code)
            print(f"Aliyun SMS response: {response}")

            # 存储验证码到 session
            request.session['sms_code'] = code
            request.session['sms_phone'] = phone

            return JsonResponse({"message": "验证码已发送，请查看手机"}, status=200)
        except Exception as e:
            return JsonResponse({"message": f"短信发送失败: {str(e)}"}, status=500)

    return JsonResponse({"message": "无效请求"}, status=400)

def register_view(request):
    """处理用户注册"""
    if request.method == 'POST':
        username = request.POST.get('username')
        email = request.POST.get('email')
        phone = request.POST.get('phone_number')
        password = request.POST.get('password')
        password2 = request.POST.get('password2')
        input_code = request.POST.get('sms_code')

        # 校验验证码
        session_code = request.session.get('sms_code')
        session_phone = request.session.get('sms_phone')

        if not session_code or not session_phone:
            messages.error(request, '请先获取验证码')
            return redirect('register')

        if phone != session_phone:
            messages.error(request, '注册手机号与验证码手机号不匹配')
            return redirect('register')

        if input_code != session_code:
            messages.error(request, '短信验证码错误')
            return redirect('register')

        if password != password2:
            messages.error(request, '两次输入的密码不一致')
            return redirect('register')

        if CustomUser.objects.filter(username=username).exists():
            messages.error(request, '该用户名已被注册')
            return redirect('register')

        if phone and CustomUser.objects.filter(phone_number=phone).exists():
            messages.error(request, '该手机号已被注册')
            return redirect('register')

        user = CustomUser.objects.create_user(username=username, email=email, phone_number=phone, password=password)

        request.session.pop('sms_code', None)
        request.session.pop('sms_phone', None)

        messages.success(request, '注册成功，请登录')
        return redirect('login')

    return render(request, 'register.html')

def ai_detect(pdf_path):
    """
    读取上传的 PDF 文件内容，并留出 AI 检测部分。

    Args:
        pdf_path (str): PDF 文件的路径。

    Returns:
        dict: 包含 PDF 文本及检测结果（暂留空）。
    """

    # 读取 PDF 文件内容
    with pdfplumber.open(pdf_path) as pdf:
        all_text = []
        for page in pdf.pages:
            page_text = page.extract_text()
            all_text.append(page_text)

        # 将每页的文本用换行分隔拼接
        final_text = "\n".join(all_text)

    # 检测逻辑 (留空部分由用户实现)
    result, h_ratio,m_ratio,l_ratio=ai_detect_chi(final_text)

    # 返回示例结果
    return result,h_ratio,m_ratio,l_ratio


def handle_file_upload(uploaded_file):
    """
    将用户上传的文件保存到 media/ 目录下，并返回保存后的文件路径。
    """
    # 确保文件名不会覆盖已有文件或引起安全问题，一般需要再做唯一化处理
    file_name = uploaded_file.name

    # 构建完整的文件保存路径
    file_path = os.path.join(settings.MEDIA_ROOT, file_name)

    # 将文件写入磁盘
    with open(file_path, 'wb+') as destination:
        for chunk in uploaded_file.chunks():
            destination.write(chunk)

    return file_path

@csrf_exempt
def upload_file(request):
    """
    第一步：只上传文件到服务器，不做AI检测。
    返回 file_id (或其他信息)，以便后续处理文件。
    """
    if request.method == 'POST':
        uploaded_file = request.FILES.get('uploaded_file')
        if not uploaded_file:
            return JsonResponse({'error': 'No file uploaded'}, status=400)
        # 保存文件到服务器
        saved_path = handle_file_upload(uploaded_file)

        # TODO: 你需要把 file_id 和 saved_path 的对应关系存储起来
        # 比如可以用数据库或缓存，这里只做示例省略
        # cache.set(file_id, saved_path, timeout=3600)

        return JsonResponse({
            'message': 'File uploaded successfully',
            'file_id': saved_path,             # 给前端后续处理用
            'file_name': uploaded_file.name, # 可选
            'file_path': saved_path         # 如果想调试看路径，可返回
        })
    else:
        return JsonResponse({'error': 'Invalid request method'}, status=405)

def reset_password_view(request):
    """
    单页单视图：在同一个页面里，
    1) 用户先填写手机号点"获取验证码"
    2) 用户拿到验证码后，再填写验证码、新密码，然后点"重置密码"
    """
    if request.method == 'POST':
        action = request.POST.get('action')  # 获取区分操作的标记: send_code / reset
        phone = request.POST.get('phone_number', '').strip()
        sms_code_input = request.POST.get('sms_code', '').strip()
        new_password = request.POST.get('password', '')
        new_password2 = request.POST.get('password2', '')
        print(phone,sms_code_input,new_password,new_password2)

        if action == 'send_code':
            # ========== 1) 用户点击"获取验证码" ==========
            if not phone:
                messages.error(request, '请输入手机号')
                return redirect('reset_password')

            # 检查手机号是否已注册
            user_exists = CustomUser.objects.filter(phone_number=phone).exists()
            if not user_exists:
                messages.error(request, '该手机号未注册')
                return redirect('reset_password')

            # 生成并发送验证码
            code = generate_code(6)
            try:
                response = send_sms(phone, code)
                print("Aliyun SMS response:", response)

                # 写入 session
                request.session['reset_phone'] = phone
                request.session['reset_code'] = code
                # 可以设置 session 有效期 (单位秒), 例如 5 分钟:
                request.session.set_expiry(300)

                messages.success(request, f'验证码已发送至 {phone}，请注意查收')
            except Exception as e:
                messages.error(request, f'短信发送失败: {str(e)}')

            return redirect('reset_password')

        elif action == 'reset':
            # ========== 2) 用户点击"重置密码" ==========
            # 从 session 读出后台保存的手机号和验证码
            session_phone = request.session.get('reset_phone')
            session_code = request.session.get('reset_code')

            # 先校验是否真的获取过验证码
            if not session_phone or not session_code:
                messages.error(request, '请先获取验证码')
                return redirect('reset_password')

            # 校验手机号是否一致
            if phone != session_phone:
                messages.error(request, '当前填写的手机号与获取验证码时不一致')
                return redirect('reset_password')

            # 校验验证码是否正确
            if sms_code_input != session_code:
                messages.error(request, '验证码错误')
                return redirect('reset_password')

            # 校验密码是否一致
            if new_password != new_password2:
                messages.error(request, '两次输入的密码不一致')
                return redirect('reset_password')

            # 在数据库找到对应用户，更新密码
            user = CustomUser.objects.filter(phone_number=phone).first()
            if not user:
                messages.error(request, '未找到对应用户，请重试')
                return redirect('reset_password')

            user.set_password(new_password)
            user.save()

            # 重置成功后清理 session
            request.session.pop('reset_phone', None)
            request.session.pop('reset_code', None)

            messages.success(request, '密码重置成功，请登录')
            return redirect('login')

    # GET 请求或未匹配到 action，都渲染同一模板
    return render(request, 'reset_password.html')

def verify_reset_code(request):
    """ 用户输入验证码，验证是否正确 """
    if request.method == 'POST':
        phone = request.POST.get('phone_number')
        input_code = request.POST.get('sms_code')

        # 读取 session 中存储的验证码
        session_code = request.session.get('reset_sms_code')
        session_phone = request.session.get('reset_phone')

        if not session_code or not session_phone:
            messages.error(request, '请先获取验证码')
            return redirect('forgot_password')

        if phone != session_phone:
            messages.error(request, '手机号不匹配，请重新输入')
            return redirect('forgot_password')

        if input_code != session_code:
            messages.error(request, '验证码错误，请重新输入')
            return redirect('forgot_password')

        # 验证通过，存储手机号到 session 以便后续重置密码
        request.session['verified_phone'] = phone

        return redirect('reset_password')

    return render(request, 'verify_reset_code.html')

def forgot_password_view(request):
    """ 用户访问找回密码页面，输入手机号获取验证码 """
    if request.method == 'POST':
        phone = request.POST.get('phone_number')

        if not phone:
            messages.error(request, "请输入手机号")
            return redirect('forgot_password')

        # 检查手机号是否存在
        if not User.objects.filter(phone_number=phone).exists():
            messages.error(request, "该手机号未注册")
            return redirect('forgot_password')

        # 生成验证码
        code = generate_code(6)

        try:
            # 发送验证码
            response = send_sms(phone, code)
            print(f"Aliyun SMS response: {response}")

            # 存储验证码到 session
            request.session['reset_sms_code'] = code
            request.session['reset_phone'] = phone

            messages.success(request, "验证码已发送，请查收短信")
            return redirect('verify_reset_code')

        except Exception as e:
            messages.error(request, f"短信发送失败: {str(e)}")
            return redirect('forgot_password')

    return render(request, 'forgot_password.html')

@csrf_exempt
def process_file(request):
    """
    第二步：对已上传的文件执行AI检测，返回检测结果和高亮文本等。
    需要前端发送 { "file_id": "xxx" } JSON。
    """
    if request.method == 'POST':

        # 解析前端发来的 JSON
        try:
            data = json.loads(request.body.decode('utf-8'))
        except:
            return JsonResponse({'error': 'Invalid JSON'}, status=400)

        file_id = data.get('file_id')
        file_name=data.get('file_name')
        if not file_id:
            return JsonResponse({'error': 'No file_id provided'}, status=400)

        # TODO: 读取 file_id 对应的文件路径
        # saved_path = cache.get(file_id) 或 从数据库查询
        # 如果找不到，则返回错误
        # if not saved_path:
        #     return JsonResponse({'error': 'Invalid file_id'}, status=404)

        # 假设这里直接演示：

        saved_path = file_id

        # 调用 AI 检测函数 (与原始函数相同)
        results, h_ratio, m_ratio, l_ratio = ai_detect(saved_path)
        h_ratio, m_ratio, l_ratio = round(h_ratio, 4)*100, round(m_ratio, 4)*100, round(l_ratio, 4)*100

        # 返回 JSON
        return JsonResponse({
            'message': 'File processed successfully',
            'filename': file_name,  # 或者从上传时的记录获取
            'file_path': saved_path,    # 或者不返回

            'ai_results': results,      # 带 "++" / "--" 标记的字符串
            "detectionResults": {
                "totalAI": str(m_ratio + h_ratio + l_ratio),
                "highAI": str(h_ratio),
                "mediumAI": str(m_ratio),
                "lowAI": str(l_ratio),
                "totalCharacters": str(len(results)),

                "distribution": {
                  "high": str(h_ratio),
                  "medium": str(m_ratio),
                  "low": str(l_ratio),
                  "none": str(100-m_ratio + h_ratio + l_ratio)
                },
                "fragments": [
                  {
                    "id": 1,
                    "originalContent": "这里是片段1的原文……",
                    "aiPercentage": "2.0%",
                    "aiValue": "0.85",
                    "suspicionLevel": "中度疑似"
                  },
                  {
                    "id": 2,
                    "originalContent": "这里是片段2的原文……",
                    "aiPercentage": "1.3%",
                    "aiValue": "0.92",
                    "suspicionLevel": "高度疑似"
                  }
                ]
            }
        })
    else:
        return JsonResponse({'error': 'Invalid request method'}, status=405)

def home(request):
    return render(request, 'home.html')  # 指定渲染的模板文件

def shuziren_api(files_dict, data_dict):
    # 替换为你的GPU服务器地址和端口（这里示例使用localhost:9000）
    api_url = 'http://localhost:9000/all_main'
    response = requests.post(api_url, files=files_dict, data=data_dict)

    if response.status_code == 200:
        print("调用成功，返回结果：", response.json())
    else:
        print("调用失败，状态码：", response.status_code)
        print("错误信息：", response.text)


def script_api(files_dict, data_dict):
    # 定义 API 地址
    api_url = 'http://localhost:9000/script'

    # 发送 POST 请求
    response = requests.post(api_url, files=files_dict, data=data_dict)

    # 打印返回结果
    if response.status_code == 200:
        print("调用成功，返回结果：", response.json())
        print(response.json()['speech_text'])
        return response.json()['speech_text']
    else:
        print("调用失败，状态码：", response.status_code)
        print("错误信息：", response.text)


def page1(request):
    """
    空白界面，直接渲染一个空白模板页面
    """
    return render(request, 'page1.html')

@csrf_exempt
def upload_video(request):
    """
    接收来自 GPU 服务器的上传视频，并保存到 MEDIA_ROOT 中
    """
    if request.method == 'POST':
        video_file = request.FILES.get('video_file')
        if not video_file:
            return JsonResponse({"status": "error", "message": "没有上传视频文件"}, status=400)
        # 保存视频到指定路径，比如 MEDIA_ROOT/videos/
        save_dir = os.path.join(settings.MEDIA_ROOT, "videos")
        if not os.path.exists(save_dir):
            os.makedirs(save_dir)
        save_path = os.path.join(save_dir, video_file.name)
        with open(save_path, 'wb+') as destination:
            for chunk in video_file.chunks():
                destination.write(chunk)

        download_url = settings.MEDIA_URL + "videos/" + video_file.name
        return JsonResponse({"status": "success", "download_url": download_url})
    else:
        return JsonResponse({"status": "error", "message": "只接受 POST 请求"}, status=405)


@login_required
def download_script(request):
    # 从 GET 参数中获取 speech_path
    speech_path = request.GET.get('speech_path')
    if not speech_path:
        return HttpResponseBadRequest("缺少下载链接参数")
    # 直接重定向到 speech_path，这里 speech_path 应该是一个有效的URL
    return redirect(speech_path)

@login_required
def generate_script(request):
    context = {}
    if request.method == 'POST':
        errors = []
        
        pdf_file = request.FILES.get('pdf_file')

        if not pdf_file:
            errors.append("请上传 PPT/PDF 文件")

        if errors:
            context['errors'] = errors
        else:
            # 构建保存目录，例如 MEDIA_ROOT/<user_id>/
            user_folder = os.path.join(settings.MEDIA_ROOT, str(request.user.id))
            if not os.path.exists(user_folder):
                os.makedirs(user_folder)
            upload_results = {}

            # 保存 PPT/PDF 文件
            try:
                pdf_path = os.path.join(user_folder, pdf_file.name)
                with open(pdf_path, 'wb+') as destination:
                    for chunk in pdf_file.chunks():
                        destination.write(chunk)
                upload_results['pdf'] = pdf_file.name
            except Exception as e:
                errors.append("PPT/PDF 文件保存失败：" + str(e))

            if errors:
                context['errors'] = errors
            else:
                context['upload_results'] = upload_results
                # 使用 PDF 文件名作为课程名称（去除 .pdf 扩展名）
                course_name = pdf_file.name
                if course_name.lower().endswith('.pdf'):
                    course_name = course_name[:-4]
                context['user_text'] = course_name

                context['upload_results'] = upload_results
                # 使用 PDF 文件名作为课程名称（去除 .pdf 扩展名）
                course_name = pdf_file.name
                if course_name.lower().endswith('.pdf'):
                    course_name = course_name[:-4]
                context['user_text'] = course_name

                files = {
                    'slides': open(pdf_path, 'rb'),  # PPT的pdf路径
                }
                data = {
                    'course_name': course_name,
                }
                response_text=script_api(files, data)

                # 创建一个新的 Word 文档
                doc = Document()

                # 添加段落
                doc.add_paragraph(response_text)

                # 保存到文件
                doc.save(pdf_path[:-4]+'.docx')
                doc_path=pdf_path[:-4] + '.docx'
                context['speech_path']=os.path.join(settings.MEDIA_URL, str(request.user.id), pdf_file.name[:-4]+'.docx')
        # 如果是 AJAX 请求，返回 JSON 数据

    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        if 'errors' in context:
            return JsonResponse({'status': 'error', 'error': context['errors']})
        else:
            return JsonResponse({'status': 'success', 'speech_path': context.get('speech_path', '')})
    # 否则返回正常渲染页面
    return render(request, 'page2.html', context)

def get_unique_filepath(directory, filename):
    """
    返回一个在指定目录下不存在的文件路径。
    如果 directory/filename 已存在，则在文件名后添加 "_new"，
    如此反复，直到找到一个不存在的文件路径。
    """
    base, ext = os.path.splitext(filename)
    candidate = os.path.join(directory, filename)
    # 记录原始 base 以便不断追加 _new
    original_base = base
    while os.path.exists(candidate):
        base = original_base + "_new" * ((len(base) - len(original_base)) // 4 + 1)
        candidate = os.path.join(directory, base + ext)
    return candidate


@login_required
def generate_course(request):
    context = {}
    if request.method == 'POST':
        errors = []

        # 获取上传的文件对象：PPT/PDF 必传，讲稿可选
        pdf_file = request.FILES.get('pdf_file')
        script_file = request.FILES.get('script_file')  # 可选上传讲稿文件
        mp4_obj = request.FILES.get('mp4_file')
        mp4_obj_already = request.POST.get('avatar_mp4_path', '')
        script_path=None
        # 构建保存目录，例如 MEDIA_ROOT/<user_id>/
        user_folder = os.path.join(settings.MEDIA_ROOT, str(request.user.id))
        if not os.path.exists(user_folder):
            os.makedirs(user_folder)
        
        # 保存 PPT/PDF 文件
        pdf_path=''
        if pdf_file:
            try:
                pdf_path = get_unique_filepath(user_folder, pdf_file.name)
                # pdf_path = os.path.join(user_folder, pdf_file.name)
                with open(pdf_path, 'wb+') as destination:
                    for chunk in pdf_file.chunks():
                        destination.write(chunk)
                course_name = pdf_path[len(user_folder)+1:]
            except Exception as e:
                errors.append("PPT/PDF 文件保存失败：" + str(e))

        # 保存讲稿文件（如果上传了）
        if script_file:
            try:
                script_path = get_unique_filepath(user_folder, script_file.name)
                # script_path = os.path.join(user_folder, script_file.name)
                with open(script_path, 'wb+') as destination:
                    for chunk in script_file.chunks():
                        destination.write(chunk)
            except Exception as e:
                errors.append("讲稿文件保存失败：" + str(e))

        mp4_path_middle=''   #不一定有，先设为空字符串
        if mp4_obj:
            try:
                mp4_path = os.path.join(user_folder, mp4_obj.name)
                mp4_path_middle=mp4_path
                # 保存 MP4 文件
                with open(mp4_path, 'wb+') as destination:
                    for chunk in mp4_obj.chunks():
                        destination.write(chunk)
            except Exception as e:
                errors.append("视频文件保存失败：" + str(e))
        else:
            mp4_path=mp4_obj_already


        if errors:
            context['errors'] = errors

        else:

            if course_name.lower().endswith('.pdf'):
                course_name = course_name[:-4]

            context['user_text'] = course_name

            # 保存一条记录到 CourseUpload 表中
            # 初始 finish 状态为 False
            CourseUpload.objects.create(
                user=request.user,
                course_name=course_name,
                mp4_file_path=mp4_path_middle,  # 此处不涉及 MP4，所以留空或其他默认值
                pdf_file_path=pdf_path,
                finish=False
            )

            if pdf_file:
                files = {
                    'slides': open(str(pdf_path), 'rb'),  # PPT的pdf路径
                    'video': open(str(mp4_path), 'rb'),     # 用户上传视频的路径
                    'speech_path':''
                }
                if script_path:
                    files['speech_path'] = open(str(script_path), 'rb')
            else:
                files = {
                    'slides': '',  # PPT的pdf路径
                    'video': open(str(mp4_path), 'rb'),     # 用户上传视频的路径
                    'speech_path':open(str(script_path), 'rb')
                }
            data = {
                'course_name': course_name,
                'output_filename': course_name + '.mp4',
                "django_upload_url": 'http://www.morecocoagents.com/upload_video/'  # 新增字段：Django上传接口的URL
            }
            shuziren_api(files, data)

    return render(request, 'page2.html', context)

@csrf_exempt
def wechat_notify(request):
    data = wxpay.callback(request.headers, request.body)
    if data and data.get("trade_state") == "SUCCESS":
        out_trade_no = data["out_trade_no"]
        Payment.objects.filter(out_trade_no=out_trade_no).update(paid_at=now())
    return JsonResponse({"code": "SUCCESS", "message": "OK"})

@login_required
def pay_discription(request):
    # 构造一个 list，方便模板遍历
    plans = []
    for key, cents in PLAN_PRICES.items():
        plans.append({
            "key": key,
            "name": PLAN_NAMES.get(key, key),
            "price": cents / 100.0,  # 单位：元
        })
    return render(request, "pay_discription.html", {"plans": plans})


@login_required
def query_order_status(request):
    out_trade_no = request.GET.get("order_id")
    mchid = settings.WECHAT_PAY["mchid"]

    print(f"[QUERY] mode={wxpay._partner_mode} mchid={mchid} out_trade_no={out_trade_no}")

    try:
        code, msg = wxpay.query(
        out_trade_no=out_trade_no
        )
        print(f"[QUERY] resp code={code} msg={msg}")

        if code != 200:
            return JsonResponse({"error": f"wechat code={code}, msg={msg}"}, status=400)

        return JsonResponse(json.loads(msg))
    except Exception as e:
        import traceback, sys
        tb = "".join(traceback.format_exception(*sys.exc_info()))
        print(tb)              # 保留日志
        return JsonResponse({"error": str(e), "trace": tb}, status=500)


@login_required
def create_native_order(request):
    # 从 GET 参数里取 plan
    plan = request.GET.get("plan", "month")
    if plan not in PLAN_PRICES:
        return HttpResponseBadRequest("无效的订阅套餐")

    user = request.user
    amount = PLAN_PRICES[plan]
    out_trade_no = uuid.uuid4().hex

    # 1. 下单到微信
    code, msg = wxpay.pay(
        description=f"{plan}订阅",
        out_trade_no=out_trade_no,
        amount={"total": amount},
        pay_type=WeChatPayType.NATIVE
    )
    if code != 200:
        return HttpResponseBadRequest(f"微信下单失败：{msg}")

    # 2. 成功后写本地库
    Payment.objects.create(
        user=user,
        out_trade_no=out_trade_no,
        amount=amount,
        paid_at=now()  # 可先写为 now()，回调时再更新真实时间
    )

    # 3. 生成二维码并渲染
    code_url = json.loads(msg)["code_url"]
    buf = io.BytesIO()
    qrcode.make(code_url).save(buf, format="PNG")
    qr_b64 = base64.b64encode(buf.getvalue()).decode()

    return render(request, "pay.html", {
        "qr_img": f"data:image/png;base64,{qr_b64}",
        "order_id": out_trade_no
    })


@login_required
def check_status(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            course_name = data.get('course_name', '').strip()
        except Exception as e:
            return HttpResponseBadRequest("JSON parse error")

        if not course_name:
            return JsonResponse({'status': 'error', 'error': '课程名称为空'})

        try:
            # 根据用户和课程名称找到对应记录
            course_upload = CourseUpload.objects.filter(user=request.user, course_name=course_name).order_by(
                '-uploaded_at').first()
            if not course_upload:
                return JsonResponse({'status': 'error', 'error': '记录不存在'})

        except CourseUpload.DoesNotExist:
            return JsonResponse({'status': 'error', 'error': '记录不存在'})

        # 构造视频文件的存储路径（假设视频文件保存在 MEDIA_ROOT/videos/<course_name>.mp4）
        video_file_path = os.path.join(settings.MEDIA_ROOT, 'videos', course_name + '.mp4')
        if os.path.exists(video_file_path):
            # 如果文件存在，更新 finish 字段为 True
            course_upload.finish = True
            course_upload.save()
            download_url = settings.MEDIA_URL + 'videos/' + course_name + '.mp4'
            return JsonResponse({'status': 'success', 'download_url': download_url})
        else:
            return JsonResponse({'status': 'pending', 'error': '视频文件尚未生成'})
    return HttpResponseBadRequest("仅支持 POST 请求")

@login_required
def get_tasks(request):
    """
    返回当前用户所有数字人生成任务请求，
    自动检查是否有视频生成完毕并更新 finish 字段，
    仅返回 course_name 非空的记录。
    """
    tasks = CourseUpload.objects.filter(user=request.user).exclude(course_name='').order_by('-uploaded_at')
    data = []

    for task in tasks:
        # 构造视频文件路径：MEDIA_ROOT/videos/<course_name>.mp4
        if task.course_name:
            video_file_path = os.path.join(settings.MEDIA_ROOT, 'videos', f"{task.course_name}.mp4")
            if os.path.exists(video_file_path) and not task.finish:
                task.finish = True
                task.save()
        local_time = timezone.localtime(task.uploaded_at).strftime('%Y-%m-%d %H:%M:%S')
        data.append({
            'id': task.id,
            'course_name': task.course_name,
            'mp4_file_path': task.mp4_file_path,
            'pdf_file_path': task.pdf_file_path,
            'finish': task.finish,
            'uploaded_at': local_time
        })

    return JsonResponse({'status': 'success', 'tasks': data})
    
@login_required
def delete_task(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            task_id = data.get('task_id')
        except Exception as e:
            return HttpResponseBadRequest("JSON解析失败")
        if not task_id:
            return JsonResponse({'status': 'error', 'error': '缺少任务ID'})
        
        try:
            task = CourseUpload.objects.get(id=task_id, user=request.user)
            task.delete()
            return JsonResponse({'status': 'success'})
        except CourseUpload.DoesNotExist:
            return JsonResponse({'status': 'error', 'error': '任务不存在'})
    return HttpResponseBadRequest("仅支持POST请求")

@login_required
def page2(request):
    context = {}
    if request.method == 'POST' and 'process' in request.POST:
        errors = []
        # 获取文本输入（虽然这里没用到）
        user_text = request.POST.get('user_text', '').strip()
        # 获取文件上传
        mp4_obj = request.FILES.get('mp4_file')

        # 校验是否上传了 MP4 文件
        if not mp4_obj:
            errors.append("请上传 MP4 视频文件")
            context['errors'] = errors
            return render(request, 'page2.html', context)

        # 构建保存目录，例如 MEDIA_ROOT/<user_id>/local_files/
        user_folder = os.path.join(settings.MEDIA_ROOT, str(request.user.id))
        os.makedirs(user_folder, exist_ok=True)

        # 构建文件保存路径
        mp4_file = os.path.join(user_folder, mp4_obj.name)

        # 保存 MP4 文件
        with open(mp4_file, 'wb+') as destination:
            for chunk in mp4_obj.chunks():
                destination.write(chunk)
                
        # 保存记录到数据库
        CourseUpload.objects.create(
            user=request.user,
            course_name='',
            mp4_file_path=mp4_file,
            pdf_file_path='',
            finish=False
        )
        
        # 添加成功提示信息
        context['message'] = "数字人保存成功！"
        return redirect('page2')  # 'page2' 为你的 URL 名称
        
    return render(request, 'page2.html', context)


@login_required
def get_finished_tasks(request):
    """
    返回当前用户满足如下条件的任务：
    1. finish=True 且 mp4_file_path 非空
    2. 或者 course_name 为空字符串
    按上传时间倒序排列，并返回课程名称、MP4 文件名及上传时间
    """
    tasks = CourseUpload.objects.filter(
        user=request.user
    ).filter(
        Q(finish=True, mp4_file_path__gt='') | Q(course_name='')
    ).order_by('-uploaded_at')

    data = []
    for task in tasks:
        mp4_filename = os.path.basename(task.mp4_file_path) if task.mp4_file_path else ""
        data.append({
            'id': task.id,
            'course_name': task.course_name,
            'mp4_filename': mp4_filename,
            'mp4_file_path': task.mp4_file_path,
            'uploaded_at': task.uploaded_at.strftime('%Y-%m-%d %H:%M:%S')
        })
    return JsonResponse({'status': 'success', 'tasks': data})


@login_required
def upload_local_file(request):
    if request.method == 'POST':
        uploaded_file = request.FILES.get('file')
        if not uploaded_file:
            return JsonResponse({'status': 'error', 'error': '未上传文件'})

        # 构建保存目录，例如 MEDIA_ROOT/<user_id>/local_files/
        user_folder = os.path.join(settings.MEDIA_ROOT, str(request.user.id), 'local_files')
        os.makedirs(user_folder, exist_ok=True)

        file_path = os.path.join(user_folder, uploaded_file.name)
        with open(file_path, 'wb+') as destination:
            for chunk in uploaded_file.chunks():
                destination.write(chunk)

        # 构造文件的 URL，确保 MEDIA_URL 配置正确
        file_url = settings.MEDIA_URL + f"{request.user.id}/local_files/" + uploaded_file.name

        # 在数据库中保存文件记录
        local_file = LocalKnowledge.objects.create(
            user=request.user,
            file_name=uploaded_file.name,
            activated=False  # 默认未激活，可根据业务修改
        )

        # 在后台线程中处理文件分块和向量化
        def process_file_in_background():
            process_file_chunks_and_vectors(file_path, local_file)
        
        processing_thread = threading.Thread(target=process_file_in_background)
        processing_thread.daemon = True
        processing_thread.start()

        file_data = {
            'id': local_file.id,  # 使用数据库记录的主键
            'name': local_file.file_name,
            'url': file_url,
            'upload_time': local_file.uploaded_at.strftime('%Y-%m-%d %H:%M:%S'),
            'activated': local_file.activated,
            'processed': local_file.processed,  # 添加处理状态
        }
        return JsonResponse({'status': 'success', 'file': file_data})
    return HttpResponseBadRequest("仅支持 POST 请求")

@login_required
def list_local_files(request):
    files_list = []
    files = LocalKnowledge.objects.filter(user=request.user).order_by('-uploaded_at')
    for f in files:
        file_url = settings.MEDIA_URL + f"{request.user.id}/local_files/" + f.file_name
        files_list.append({
            'id': f.id,
            'name': f.file_name,
            'url': file_url,
            'upload_time': f.uploaded_at.strftime('%Y-%m-%d %H:%M:%S'),
            'activated': f.activated,
            'processed': f.processed,
            'processing_error': f.processing_error,
        })
    return JsonResponse({'status': 'success', 'files': files_list})


@login_required
def delete_local_file(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            file_id = data.get('file_id')
            if not file_id:
                return JsonResponse({'status': 'error', 'error': '未提供文件标识'})
        except Exception as e:
            return HttpResponseBadRequest("JSON解析错误")

        try:
            # 根据 file_id 和当前用户查找记录
            local_file = LocalKnowledge.objects.get(id=file_id, user=request.user)
        except LocalKnowledge.DoesNotExist:
            return JsonResponse({'status': 'error', 'error': '文件记录不存在'})

        # 删除磁盘上文件
        user_folder = os.path.join(settings.MEDIA_ROOT, str(request.user.id), 'local_files')
        file_path = os.path.join(user_folder, local_file.file_name)
        if os.path.exists(file_path):
            os.remove(file_path)
        # 删除数据库记录
        local_file.delete()
        return JsonResponse({'status': 'success'})
    return HttpResponseBadRequest("仅支持 POST 请求")


@login_required
def update_local_file_activation(request):
    import time
    start_time = time.time()
    print(f"[DEBUG后端] 请求开始: 时间={time.strftime('%H:%M:%S')}")
    
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            file_id = data.get('file_id')
            activated = data.get('activated', False)
            print(f"[DEBUG后端] 解析参数完成: file_id={file_id}, activated={activated}, 耗时={time.time()-start_time:.3f}s")
        except Exception:
            return HttpResponseBadRequest("JSON parse error")

        try:
            query_start = time.time()
            # 只查询必要的字段，避免加载大型JSONField
            local_file = LocalKnowledge.objects.only('id', 'activated', 'user_id').get(id=file_id, user=request.user)
            print(f"[DEBUG后端] 数据库查询完成: 耗时={time.time()-query_start:.3f}s")
        except LocalKnowledge.DoesNotExist:
            return JsonResponse({'status': 'error', 'error': '文件记录不存在'})

        save_start = time.time()
        # 使用update()方法直接更新，避免加载整个对象
        updated_count = LocalKnowledge.objects.filter(id=file_id, user=request.user).update(activated=activated)
        print(f"[DEBUG后端] 数据库保存完成: 耗时={time.time()-save_start:.3f}s, 更新记录数={updated_count}")
        
        print(f"[DEBUG后端] 请求完成: 总耗时={time.time()-start_time:.3f}s")
        return JsonResponse({'status': 'success'})
    return HttpResponseBadRequest("仅支持 POST 请求")

@login_required
def check_file_processing_status(request):
    """检查文件处理状态的API"""
    if request.method == 'GET':
        file_id = request.GET.get('file_id')
        if not file_id:
            return JsonResponse({'status': 'error', 'error': '未提供文件ID'})
            
        try:
            local_file = LocalKnowledge.objects.get(id=file_id, user=request.user)
            return JsonResponse({
                'status': 'success',
                'processed': local_file.processed,
                'processing_error': local_file.processing_error,
                'chunks_count': len(local_file.chunks) if local_file.chunks else 0
            })
        except LocalKnowledge.DoesNotExist:
            return JsonResponse({'status': 'error', 'error': '文件记录不存在'})
    return HttpResponseBadRequest("仅支持 GET 请求")

@login_required
def reprocess_file(request):
    """重新处理文件的API"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            file_id = data.get('file_id')
        except Exception:
            return HttpResponseBadRequest("JSON parse error")
            
        try:
            local_file = LocalKnowledge.objects.get(id=file_id, user=request.user)
        except LocalKnowledge.DoesNotExist:
            return JsonResponse({'status': 'error', 'error': '文件记录不存在'})
            
        # 重新处理文件
        file_path = os.path.join(settings.MEDIA_ROOT, str(request.user.id), 'local_files', local_file.file_name)
        
        if not os.path.exists(file_path):
            return JsonResponse({'status': 'error', 'error': '文件不存在'})
            
        # 在后台线程中重新处理
        def reprocess_file_in_background():
            process_file_chunks_and_vectors(file_path, local_file)
        
        processing_thread = threading.Thread(target=reprocess_file_in_background)
        processing_thread.daemon = True
        processing_thread.start()
        
        return JsonResponse({'status': 'success', 'message': '已开始重新处理文件'})
    return HttpResponseBadRequest("仅支持 POST 请求")

@login_required
def page_loca(request):
    return render(request, 'page_loca.html')

@login_required  # 如果需要登录后才能访问，可加此装饰器
def my_categories_view(request):
    """
    返回一个卡片式功能选择页面（my_categories.html）
    """
    # 如果需要在模板中传递额外上下文，可在此构造字典
    context = {}
    return render(request, 'my_categories.html', context)

def my_categories_view_2(request):
    return render(request, 'my_categories_2.html')

def my_categories_view_3(request):
    return render(request, 'my_categories_3.html')

def my_categories_view_4(request):
    return render(request, 'my_categories_4.html')



def real_home(request):
    return render(request, 'real_home.html')  # 指定渲染的模板文件

@login_required
@csrf_exempt  # 生产环境请用 CSRF Token 而非禁用
def delete_conversation(request):
    """
    删除当前用户的一个聊天会话
    请求体 JSON 格式：{"conversation_id": 1}
    """
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            conv_id = data.get('conversation_id')
            conv = Conversation.objects.get(id=conv_id, user=request.user)
            conv.delete()
            return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'error': str(e)})
    return JsonResponse({'status': 'error', 'error': 'Invalid request method'})

@login_required
def get_conversations(request):
    """
    返回当前用户所有的聊天会话及对应消息，
    按"最后活动时间"倒序排序：
      - 有消息时使用最后一条消息的时间
      - 没有消息时使用会话本身的创建时间
    """
    conversations = (
        Conversation.objects.filter(user=request.user)
        .annotate(
            last_activity=Coalesce(
                Max('messages__created_at'),  # 有消息时使用最后消息时间
                F('created_at')              # 否则使用会话创建时间
            )
        )
        .order_by('-last_activity')  # 倒序排列
    )

    data = []
    for conv in conversations:
        messages = conv.messages.all().order_by('created_at')
        msg_data = [
            {
                'sender': m.sender,
                'message': m.message,
                'attachment': m.attachment,
                'summary': m.summary,
                'chunk_data': m.chunk_data,
                'created_at': m.created_at.strftime('%Y-%m-%d %H:%M:%S')
            }
            for m in messages
        ]
        data.append({
            'id': conv.id,
            'icon': conv.icon,
            'messages': msg_data,
        })
    return JsonResponse({'conversations': data})

@login_required
@csrf_exempt  # 生产环境建议使用 CSRF Token 机制，此处为示例方便
def create_conversation(request):
    """
    创建一个新的聊天会话
    请求体 JSON 格式：{ "icon": "🤖" }
    """
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            icon = data.get('icon', '')
        except Exception:
            return HttpResponseBadRequest("JSON parse error")
        conv = Conversation.objects.create(user=request.user, icon=icon)
        return JsonResponse({'id': conv.id, 'icon': conv.icon})
    return HttpResponseBadRequest("Invalid request method")

@login_required
@csrf_exempt  # 生产环境请使用 CSRF Token 验证
def upload_chat_file(request):
    if request.method == 'POST':
        uploaded_file = request.FILES.get('file')
        conversation_id = request.POST.get('conversation_id')
        if not uploaded_file:
            return JsonResponse({'status': 'error', 'error': '未上传文件'})
        # 构建保存路径，例如：MEDIA_ROOT/<user_id>/chat_files/
        user_folder = os.path.join(settings.MEDIA_ROOT, str(request.user.id), 'chat_files')
        if not os.path.exists(user_folder):
            os.makedirs(user_folder)
        file_path = os.path.join(user_folder, uploaded_file.name)
        with open(file_path, 'wb+') as destination:
            for chunk in uploaded_file.chunks():
                destination.write(chunk)
        # # 构造文件的访问 URL（确保 MEDIA_URL 配置正确）
        # file_url = settings.MEDIA_URL + f"{request.user.id}/chat_files/" + uploaded_file.name
        return JsonResponse({'status': 'success', 'file_url': file_path})
    return JsonResponse({'status': 'error', 'error': '仅支持 POST 请求'})


def process_file_chunks_and_vectors(file_path, local_knowledge_obj, chunk_size=2000, overlap=50):
    """
    处理上传的文件，进行分块和向量化，并保存到数据库
    
    参数：
      file_path: 文件在服务器上的路径
      local_knowledge_obj: LocalKnowledge对象
      chunk_size: 分块大小
      overlap: 重叠字符数
    
    返回：
      处理成功返回True，失败返回False
    """
    try:
        # 1. 提取文件内容
        file_content = ""
        file_extension = file_path.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            with pdfplumber.open(file_path) as pdf:
                all_text = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        all_text.append(page_text)
                file_content = "\n".join(all_text)
        elif file_extension in ['docx', 'doc']:
            file_content = read_word_file(file_path)
        elif file_extension == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                file_content = f.read()
        else:
            raise ValueError(f"不支持的文件格式: {file_extension}")
            
        if not file_content.strip():
            raise ValueError("文件内容为空")
            
        # 2. 文本分块
        chunks = []
        step = chunk_size - overlap
        for i in range(0, len(file_content), step):
            chunk_text = file_content[i:i + chunk_size].strip()
            if chunk_text:
                chunks.append({
                    'text': chunk_text,
                    'index': len(chunks)
                })
                
        if not chunks:
            raise ValueError("无法生成有效的文本块")
            
        # 3. 生成向量
        embedding_fn = OpenAIEmbeddings()
        chunk_texts = [chunk['text'] for chunk in chunks]
        vectors = embedding_fn.embed_documents(chunk_texts)
        
        # 4. 更新数据库记录
        local_knowledge_obj.file_content = file_content
        local_knowledge_obj.chunks = chunks
        local_knowledge_obj.vectors = vectors
        local_knowledge_obj.chunk_size = chunk_size
        local_knowledge_obj.overlap = overlap
        local_knowledge_obj.processed = True
        local_knowledge_obj.processing_error = None
        local_knowledge_obj.save()
        
        return True
        
    except Exception as e:
        # 保存错误信息
        local_knowledge_obj.processed = False
        local_knowledge_obj.processing_error = str(e)
        local_knowledge_obj.save()
        print(f"文件处理错误: {str(e)}")
        return False

def get_relevant_text_from_vectors(local_knowledge_obj, query, top_k=4):
    """
    使用预计算的向量进行相似度搜索
    
    参数：
      local_knowledge_obj: LocalKnowledge对象
      query: 查询文本
      top_k: 返回最相关的前k个结果
    
    返回：
      相关文本块列表
    """
    if not local_knowledge_obj.processed or not local_knowledge_obj.vectors:
        return []
        
    try:
        # 1. 生成查询向量
        embedding_fn = OpenAIEmbeddings()
        query_vector = embedding_fn.embed_query(query)
        
        # 2. 构建FAISS索引
        vectors_array = np.array(local_knowledge_obj.vectors, dtype=np.float32)
        embedding_size = vectors_array.shape[1]
        index = faiss.IndexFlatL2(embedding_size)
        index.add(vectors_array)
        
        # 3. 搜索最相似的向量
        query_vector_array = np.array([query_vector], dtype=np.float32)
        distances, indices = index.search(query_vector_array, min(top_k, len(local_knowledge_obj.chunks)))
        
        # 4. 返回对应的文本块
        relevant_chunks = []
        if len(indices) > 0 and len(indices[0]) > 0:
            for idx in indices[0]:
                if idx < len(local_knowledge_obj.chunks):
                    relevant_chunks.append(local_knowledge_obj.chunks[idx]['text'])
                
        return relevant_chunks
        
    except Exception as e:
        print(f"向量搜索错误: {str(e)}")
        return []

def get_most_relevant_text(long_text, prompt, chunk_size=2000, overlap=50):
    """
    输入长文本和查询的 prompt，返回最相关的文本块。

    参数：
      long_text: 待处理的长文本字符串
      prompt: 用户查询的 prompt
      chunk_size: 每个文本块的最大字符数（默认500）
      overlap: 相邻文本块的重叠字符数（默认50）

    返回：
      与查询最相关的文本块字符串，如果没有则返回 None。
    """
    # 1. 分块：采用滑动窗口方式进行文本分块
    chunks = []
    step = chunk_size - overlap
    for i in range(0, len(long_text), step):
        chunks.append(long_text[i:i + chunk_size])

    # 2. 构建 FAISS 向量库
    embedding_size = 1536  # OpenAIEmbeddings 的维度
    index = faiss.IndexFlatL2(embedding_size)
    embedding_fn = OpenAIEmbeddings().embed_query
    vectorstore = FAISS(embedding_fn, index, InMemoryDocstore({}), {})

    # 3. 将所有文本块添加到向量库中
    vectorstore.add_texts(chunks)

    # 4. 根据查询检索最相关的文本块，取 top1
    results = vectorstore.similarity_search(prompt, k=4)
    new=[p.page_content for p in results]
    # new="\n".join(new)
    if results:
        return new
    else:
        return None

def clean_pdf_text(raw_text: str) -> str:
    """
    Fix common formatting issues in text extracted from PDF.
    """
    # 1) 把 form-feed 替换成换行，方便后续处理
    txt = raw_text.replace("\x0c", "\n")

    # 2) 拼接被 '-\n' 打断的单词（de-hyphenation）
    txt = re.sub(r"(\w+)-\n(\w+)", r"\1\2", txt)

    # 3) 若行尾不是句子结束符，\n → 空格（保持段落完整）
    txt = re.sub(r"(?<![.\?!])\n(?=\w)", " ", txt)

    # 4) 检测并删除重复出现的页眉/页脚行
    lines: List[str] = txt.splitlines()
    counts = Counter(lines)

    # 估个页数阈值：出现次数 ≥ 一半页数就视为页眉/脚
    estimated_pages = max(txt.count("\n") // 40, 1)
    repeat_threshold = max(2, estimated_pages // 2)

    cleaned_lines = [
        ln for ln in lines
        if counts[ln] < repeat_threshold
        and not re.match(r"^\s*Page\s+\d+\s*$", ln)  # 去除 'Page 3' 之类
        and ln.strip()                               # 去除全空白行
    ]
    txt = "\n".join(cleaned_lines)

    # 5) 压缩多余空格、空行
    txt = re.sub(r"[ \t]+", " ", txt)        # 多空格 → 1 空格
    txt = re.sub(r"\n{3,}", "\n\n", txt)     # 连续 ≥3 换行 → 2 换行（段落分隔）
    return txt.strip()

@login_required
@csrf_exempt
def send_message(request):
    """
    向指定会话中发送消息，并模拟机器人回复。
    """
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            conversation_id = data.get('conversation_id')
            message_text = data.get('message', '').strip()
            attachment = data.get('attachment', '').strip()
            local_knowledge = data.get('local_knowledge', False)
            chat_model = data.get('chat_model', '')  # 获取聊天模型选项
        except Exception:
            return HttpResponseBadRequest("JSON parse error")
        if not message_text:
            return HttpResponseBadRequest("Empty message")
        try:
            conv = Conversation.objects.get(id=conversation_id, user=request.user)
        except Conversation.DoesNotExist:
            return HttpResponseBadRequest("Conversation not found")
        icon = conv.icon

        # 提取该会话所有的聊天记录，按创建时间升序排序
        messages_old = conv.messages.all().order_by('created_at')
        messages_data = [{
            'sender': m.sender,
            'real_message': m.message,
        } for m in messages_old]
        chat_history=convert_messages(messages_data)
        this_img=[]
        message_text1 = message_text
        if attachment[-3:]=='png' or attachment[-3:]=='jpg':
            this_img.append(attachment)
        if attachment[-3:]=='pdf':
            with pdfplumber.open(attachment) as pdf:
                all_text = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    all_text.append(page_text)
                # 将每页的文本用换行分隔拼接
                final_text = "\n".join(all_text)
            raged_text_list=get_most_relevant_text(final_text, message_text1)
            raged_text=''.join(raged_text_list)
            message_text1=message_text1+raged_text
        if attachment[-4:]=='docx':
            all_text=read_word_file(attachment)
            raged_text_list = get_most_relevant_text(all_text, message_text1)
            raged_text = ''.join(raged_text_list)
            message_text1 = message_text1 + raged_text
        summary=None
        if conv.messages.count() == 0:
            summary=get_response_with_img('请你总结一下用户的问题，总结在10个字以内', message_text1, [], this_img)
        if local_knowledge:
            from .models import LocalKnowledge
            activated_files = LocalKnowledge.objects.filter(user=request.user, activated=True, processed=True)
            
            if not activated_files.exists():
                robot_answer = "暂无已激活且处理完成的本地知识库文件，请先上传并激活文件。"
                all_raged_dict = {}
            else:
                print(f"使用本地知识库，已激活文件数量：{activated_files.count()}")
            
            all_raged_text=[]
            all_raged_dict={}
                
            for i, lk in enumerate(activated_files):
                # 使用预计算的向量进行搜索
                raged_text = get_relevant_text_from_vectors(lk, message_text1, top_k=4)
                
                for j, chunk_text in enumerate(raged_text):
                    chunk_key = f'doc{i}#chunk{j}'
                    chunk_text=clean_pdf_text(chunk_text)
                    all_raged_dict[chunk_key] = chunk_text
                    formatted_chunk = f'[{chunk_key}] ' + chunk_text
                    all_raged_text.append(formatted_chunk)
            
            if all_raged_text:
                notebook_output = my_notebook(message_text1, all_raged_text)
                robot_answer = generate_markdown(notebook_output)
            else:
                robot_answer = "未找到相关的知识库内容。"
                all_raged_dict = {}
        else:
            all_raged_dict = {}
            # 当不使用本地知识库时，使用普通的AI回复
            robot_answer = get_response_with_img('你是一个智能助手', message_text1, chat_history, this_img)
            
         # 保存用户消息
        ChatMessage.objects.create(conversation=conv, sender='user', message=message_text, real_messages=message_text1, summary=summary if summary else None, attachment=attachment if attachment else None)
        ChatMessage.objects.create(conversation=conv, sender='bot', message=robot_answer, chunk_data=all_raged_dict if all_raged_dict else None)

        return JsonResponse({'status': 'success'})
    return HttpResponseBadRequest("Invalid request method")


@login_required
@csrf_exempt
def send_message_stream(request):
    """
    流式发送消息端点，使用Server-Sent Events实现真正的流式输出
    """
    if request.method != 'POST':
        return HttpResponseBadRequest("Invalid request method")
    
    try:
        data = json.loads(request.body)
        conversation_id = data.get('conversation_id')
        message_text = data.get('message', '').strip()
        attachment = data.get('attachment', '').strip()
        local_knowledge = data.get('local_knowledge', False)
        chat_model = data.get('chat_model', '')
    except Exception:
        return HttpResponseBadRequest("JSON parse error")
    
    if not message_text:
        return HttpResponseBadRequest("Empty message")
    
    try:
        conv = Conversation.objects.get(id=conversation_id, user=request.user)
    except Conversation.DoesNotExist:
        return HttpResponseBadRequest("Conversation not found")
    
    def generate_response():
        try:
            # 处理消息和附件（与原版相同）
            messages_old = conv.messages.all().order_by('created_at')
            messages_data = [{
                'sender': m.sender,
                'real_message': m.message,
            } for m in messages_old]
            chat_history = convert_messages(messages_data)
            this_img = []
            message_text1 = message_text
            
            # 处理附件
            if attachment:
                if attachment[-3:] in ['png', 'jpg']:
                    this_img.append(attachment)
                elif attachment[-3:] == 'pdf':
                    with pdfplumber.open(attachment) as pdf:
                        all_text = []
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            all_text.append(page_text)
                        final_text = "\n".join(all_text)
                    raged_text_list = get_most_relevant_text(final_text, message_text1)
                    raged_text = ''.join(raged_text_list)
                    message_text1 = message_text1 + raged_text
                elif attachment[-4:] == 'docx':
                    all_text = read_word_file(attachment)
                    raged_text_list = get_most_relevant_text(all_text, message_text1)
                    raged_text = ''.join(raged_text_list)
                    message_text1 = message_text1 + raged_text
            
            # 生成总结（如果是首次对话）
            summary = None
            if conv.messages.count() == 0:
                summary = get_response_with_img('请你总结一下用户的问题，总结在10个字以内', message_text1, [], this_img)
            
            # 先保存用户消息
            user_message = ChatMessage.objects.create(
                conversation=conv, 
                sender='user', 
                message=message_text, 
                real_messages=message_text1, 
                summary=summary if summary else None, 
                attachment=attachment if attachment else None
            )
            
            # 发送消息保存确认
            yield f"data: {json.dumps({'type': 'user_message_saved', 'message_id': user_message.id})}\n\n"
            
            # 处理本地知识库
            all_raged_dict = {}
            if local_knowledge:
                from .models import LocalKnowledge
                activated_files = LocalKnowledge.objects.filter(user=request.user, activated=True, processed=True)
                
                if not activated_files.exists():
                    # 直接返回错误消息
                    error_msg = "暂无已激活且处理完成的本地知识库文件，请先上传并激活文件。"
                    yield f"data: {json.dumps({'type': 'token', 'content': error_msg})}\n\n"
                    yield f"data: {json.dumps({'type': 'done'})}\n\n"
                    
                    # 保存机器人消息
                    ChatMessage.objects.create(conversation=conv, sender='bot', message=error_msg)
                    return
                
                # 处理知识库文本
                all_raged_text = []
                for i, lk in enumerate(activated_files):
                    raged_text = get_relevant_text_from_vectors(lk, message_text1, top_k=4)
                    for j, chunk_text in enumerate(raged_text):
                        chunk_key = f'doc{i}#chunk{j}'
                        chunk_text = clean_pdf_text(chunk_text)
                        all_raged_dict[chunk_key] = chunk_text
                        formatted_chunk = f'[{chunk_key}] ' + chunk_text
                        all_raged_text.append(formatted_chunk)
                
                if all_raged_text:
                    # 使用流式版本的notebook处理
                    from .my_notebook import my_notebook_stream
                    
                    # 流式输出markdown格式的tokens
                    full_response = ""
                    for token in my_notebook_stream(message_text1, all_raged_text):
                        full_response += token
                        yield f"data: {json.dumps({'type': 'token', 'content': token})}\n\n"
                    
                    yield f"data: {json.dumps({'type': 'done', 'chunk_data': all_raged_dict})}\n\n"
                    
                    # 保存机器人消息
                    ChatMessage.objects.create(conversation=conv, sender='bot', message=full_response, chunk_data=all_raged_dict)
                else:
                    error_msg = "未找到相关的知识库内容。"
                    yield f"data: {json.dumps({'type': 'token', 'content': error_msg})}\n\n"
                    yield f"data: {json.dumps({'type': 'done'})}\n\n"
                    
                    ChatMessage.objects.create(conversation=conv, sender='bot', message=error_msg)

            else:
                # 普通AI回复 - 使用流式API
                yield from stream_openai_response(message_text1, chat_history, this_img, conv)
                
        except (json.JSONDecodeError, ValueError, TypeError, AttributeError) as e:
            # 捕获特定的常见错误
            error_msg = f"数据处理错误: {str(e)}"
            yield f"data: {json.dumps({'type': 'error', 'message': error_msg})}\n\n"
        except Exception as e:
            # 其他未预期的错误，显示详细信息用于调试
            import traceback
            error_msg = f"系统错误: {str(e)}\n{traceback.format_exc()}"
            yield f"data: {json.dumps({'type': 'error', 'message': error_msg})}\n\n"
    
    def stream_generator():
        # 立即发送一个初始消息确保连接建立
        yield f"data: {json.dumps({'type': 'connected'})}\n\n"
        # 然后开始真正的响应生成
        yield from generate_response()

    response = HttpResponse(stream_generator(), content_type='text/event-stream; charset=utf-8')
    response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response['Pragma'] = 'no-cache'
    response['Expires'] = '0'
    response['Access-Control-Allow-Origin'] = '*'
    response['Access-Control-Allow-Headers'] = 'Cache-Control'
    response['X-Accel-Buffering'] = 'no'  # 禁用Nginx缓冲
    
    return response


def stream_openai_response(message_text, chat_history, image_list, conversation):
    """
    流式调用OpenAI API并逐个token返回
    """
    try:
        full_response = ""
        # 使用新的流式函数
        for token in get_response_with_img_stream('你是一个智能助手', message_text, chat_history, image_list):
            full_response += token
            yield f"data: {json.dumps({'type': 'token', 'content': token})}\n\n"
        
        # 发送完成信号
        yield f"data: {json.dumps({'type': 'done'})}\n\n"
        
        # 保存完整的机器人回复到数据库
        ChatMessage.objects.create(conversation=conversation, sender='bot', message=full_response)
        
    except (ConnectionError, TimeoutError) as e:
        error_msg = f"网络连接错误: {str(e)}"
        yield f"data: {json.dumps({'type': 'error', 'message': error_msg})}\n\n"
    except Exception as e:
        import traceback
        error_msg = f"API调用错误: {str(e)}\n{traceback.format_exc()}"
        yield f"data: {json.dumps({'type': 'error', 'message': error_msg})}\n\n"